"""
Document loader with structure-aware chunking and table extraction.
Splits on document structure (sections, paragraphs) instead of raw character counts.
"""
from __future__ import annotations

import os
import re
import hashlib
import csv
import fitz  # pymupdf


def compute_file_hash(file_path: str) -> str:
    """Compute SHA256 hash of a file (for deduplication)."""
    sha = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            sha.update(chunk)
    return sha.hexdigest()


def load_pdf_text(file_path: str) -> str:
    """Extract full text from a PDF with page markers. Returns empty string on failure."""
    try:
        doc = fitz.open(file_path)
        text = ""
        for i, page in enumerate(doc, 1):
            page_text = page.get_text()
            if page_text.strip():
                text += f"\n[PAGE_{i}]\n{page_text}"
        doc.close()
        return text
    except Exception as e:
        print(f"  WARN: Failed to parse {file_path}: {e}")
        return ""


def load_docx_text(file_path: str) -> str:
    """Extract paragraphs and tables from a DOCX file."""
    try:
        from docx import Document
    except ImportError as e:
        raise ImportError("python-docx is required for DOCX support. Install with: pip install python-docx") from e

    try:
        doc = Document(file_path)
        parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        for table_index, table in enumerate(doc.tables, 1):
            rows = []
            for row in table.rows:
                cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
                if any(cells):
                    rows.append(cells)
            if rows:
                parts.append(f"[TABLE_{table_index}]\n{_table_rows_to_markdown(rows)}")

        return "\n\n".join(parts)
    except Exception as e:
        print(f"  WARN: Failed to parse {file_path}: {e}")
        return ""


def load_text_file(file_path: str) -> str:
    """Load TXT/Markdown/CSV text with a few common encodings."""
    suffix = os.path.splitext(file_path)[1].lower()
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            if suffix == ".csv":
                return _load_csv_as_markdown(file_path, encoding)
            with open(file_path, "r", encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
        except Exception as e:
            print(f"  WARN: Failed to parse {file_path}: {e}")
            return ""
    print(f"  WARN: Failed to decode {file_path}")
    return ""


def load_document_text(file_path: str) -> str:
    """Extract text from a supported document path."""
    suffix = os.path.splitext(file_path)[1].lower()
    if suffix == ".pdf":
        return load_pdf_text(file_path)
    if suffix == ".docx":
        return load_docx_text(file_path)
    if suffix in (".txt", ".md", ".markdown", ".csv"):
        return load_text_file(file_path)
    print(f"  WARN: Unsupported document type: {file_path}")
    return ""


def extract_tables_from_pdf(file_path: str) -> list[dict]:
    """
    Extract tables from a PDF and convert to Markdown format.

    Uses PyMuPDF's built-in table detection (fitz.Page.find_tables).
    Returns: [{'page': int, 'markdown': str, 'rows': int, 'cols': int}]

    Note: find_tables() requires PyMuPDF >= 1.23.0.
    Falls back gracefully on older versions.
    """
    tables = []
    try:
        doc = fitz.open(file_path)
        for page_num, page in enumerate(doc, 1):
            try:
                found = page.find_tables()
            except AttributeError:
                # find_tables not available (PyMuPDF < 1.23.0)
                doc.close()
                return tables

            for table in found:
                data = table.extract()
                if not data or len(data) < 2:
                    continue

                # Convert to Markdown table
                rows = len(data)
                cols = max(len(row) for row in data) if data else 0
                md_lines = []

                # Header row
                header = [str(cell or "").replace("\n", " ").strip() for cell in data[0]]
                # Pad header to match max cols
                while len(header) < cols:
                    header.append("")
                md_lines.append("| " + " | ".join(header) + " |")
                md_lines.append("| " + " | ".join(["---"] * cols) + " |")

                # Data rows
                for row in data[1:]:
                    cells = [str(cell or "").replace("\n", " ").strip() for cell in row]
                    while len(cells) < cols:
                        cells.append("")
                    md_lines.append("| " + " | ".join(cells) + " |")

                tables.append({
                    "page": page_num,
                    "markdown": "\n".join(md_lines),
                    "rows": rows,
                    "cols": cols,
                })
        doc.close()
    except Exception as e:
        print(f"  WARN: Table extraction failed for {file_path}: {e}")

    return tables


def _table_rows_to_markdown(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    cols = max(len(row) for row in rows)
    normalized = []
    for row in rows:
        cells = [str(cell or "").replace("\n", " ").strip() for cell in row]
        while len(cells) < cols:
            cells.append("")
        normalized.append(cells)
    lines = ["| " + " | ".join(normalized[0]) + " |"]
    lines.append("| " + " | ".join(["---"] * cols) + " |")
    for row in normalized[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def _load_csv_as_markdown(file_path: str, encoding: str) -> str:
    rows: list[list[str]] = []
    with open(file_path, "r", encoding=encoding, newline="") as f:
        reader = csv.reader(f)
        for idx, row in enumerate(reader):
            rows.append(row)
            if idx >= 500:
                rows.append(["..."])
                break
    return _table_rows_to_markdown(rows)


def _extract_year(meta: dict) -> str:
    """Extract year from PDF metadata (various date formats)."""
    for key in ("creationDate", "modDate", "date"):
        val = meta.get(key, "")
        if val:
            m = re.search(r"(?:D:)?(\d{4})", str(val))
            if m:
                return m.group(1)
    return ""


def load_pdf_meta(file_path: str) -> dict:
    """
    Extract metadata (title, author, year, file_hash) from a PDF.
    Falls back to filename if metadata is missing.
    """
    try:
        doc = fitz.open(file_path)
        meta = doc.metadata
        doc.close()

        title = (meta.get("title") or "").strip()
        author = (meta.get("author") or "").strip()
        year = _extract_year(meta)
        file_hash = compute_file_hash(file_path)

        return {
            "title": title,
            "author": author,
            "year": year,
            "file_hash": file_hash,
        }
    except Exception:
        return {"title": "", "author": "", "year": "", "file_hash": ""}


def load_document_meta(file_path: str) -> dict:
    """Extract best-effort metadata for any supported document."""
    suffix = os.path.splitext(file_path)[1].lower()
    if suffix == ".pdf":
        return load_pdf_meta(file_path)
    meta = {
        "title": os.path.splitext(os.path.basename(file_path))[0],
        "author": "",
        "year": "",
        "file_hash": compute_file_hash(file_path),
    }
    if suffix == ".docx":
        try:
            from docx import Document
            props = Document(file_path).core_properties
            meta["title"] = props.title or meta["title"]
            meta["author"] = props.author or ""
            if props.created:
                meta["year"] = str(props.created.year)
        except Exception:
            pass
    return meta


def load_pdfs_from_dir(dir_path: str) -> list[dict]:
    """Load all PDFs from a directory. Returns [{'name': ..., 'text': ...}]."""
    documents = []
    for fname in os.listdir(dir_path):
        if fname.endswith(".pdf"):
            full_path = os.path.join(dir_path, fname)
            text = load_pdf_text(full_path)
            if text.strip():
                documents.append({"name": fname, "text": text, "path": full_path})
    return documents


def load_documents_from_dir(dir_path: str) -> list[dict]:
    """Load all supported documents from a directory."""
    documents = []
    for fname in os.listdir(dir_path):
        suffix = os.path.splitext(fname)[1].lower()
        if suffix in (".pdf", ".docx", ".txt", ".md", ".markdown", ".csv"):
            full_path = os.path.join(dir_path, fname)
            text = load_document_text(full_path)
            if text.strip():
                documents.append({"name": fname, "text": text, "path": full_path})
    return documents


# ── Structure-aware chunking ─────────────────────────────────

_PAGE_RE = re.compile(r"\[PAGE_(\d+)\]")


def _extract_pages(text: str) -> str:
    """Extract page numbers from [PAGE_N] markers."""
    pages = _PAGE_RE.findall(text)
    return ",".join(pages) if pages else "?"


def _strip_page_markers(text: str) -> str:
    """Remove [PAGE_N] markers from text."""
    return _PAGE_RE.sub("", text).strip()


SECTION_PATTERNS = re.compile(
    r"\n\s*(?:\d+[\.\)]\s*)?"
    r"(?:Abstract|Introduction|Related\s*Work|Background|"
    r"Method|Methodology|Data|Empirical\s*Strategy|Results?|"
    r"Discussion|Conclusion|References?|Bibliography|Appendix|"
    r"Motivation|Model|Experiment|Evaluation|Findings|Policy\s*Implications)",
    re.IGNORECASE,
)


def split_on_sections(text: str) -> list[str]:
    """
    Split text on common section headers.
    Sections that are too long are further split by paragraphs.
    """
    matches = list(SECTION_PATTERNS.finditer(text))
    if len(matches) < 2:
        return _split_paragraphs(text)

    sections = []
    for i, match in enumerate(matches):
        start = match.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        section_text = text[start:end].strip()
        if len(section_text) > 100:
            sections.append(section_text)

    return sections


def _split_paragraphs(text: str, max_chars: int = 1200) -> list[str]:
    """Split text by paragraph boundaries, merging short ones."""
    paragraphs = re.split(r"\n\s*\n", text)
    chunks = []
    current = ""
    for p in paragraphs:
        p = p.strip()
        if not p:
            continue
        if len(current) + len(p) < max_chars:
            current += ("\n\n" if current else "") + p
        else:
            if current:
                chunks.append(current)
            current = p
    if current:
        chunks.append(current)
    return chunks


def chunk_documents(
    documents: list[dict],
    extract_tables: bool = True,
) -> list[dict]:
    """
    Structure-aware chunking pipeline with optional table extraction:
    1. Split each document into sections
    2. Split oversized sections into paragraphs
    3. Optionally extract tables and append as separate chunks

    documents: [{'name': ..., 'text': ..., 'path'?: ...}]
    Returns: [{'source_name': ..., 'chunk_id': ..., 'text': ..., 'page': ...,
               'is_table': bool}]
    """
    chunks = []
    chunk_id = 0

    for doc in documents:
        sections = split_on_sections(doc["text"])

        if not sections or len(sections) == 1:
            sections = _split_paragraphs(doc["text"])

        for section in sections:
            if len(section) > 2000:
                sub = _split_paragraphs(section, max_chars=1000)
                for sub_text in sub:
                    if sub_text.strip() and len(sub_text) > 50:
                        chunks.append({
                            "source_name": doc["name"],
                            "chunk_id": chunk_id,
                            "text": _strip_page_markers(sub_text.strip()),
                            "page": _extract_pages(sub_text.strip()),
                            "is_table": False,
                        })
                        chunk_id += 1
            else:
                if section.strip() and len(section) > 50:
                    chunks.append({
                        "source_name": doc["name"],
                        "chunk_id": chunk_id,
                        "text": _strip_page_markers(section.strip()),
                        "page": _extract_pages(section.strip()),
                        "is_table": False,
                    })
                    chunk_id += 1

        # ── PDF table extraction ──
        if extract_tables and "path" in doc and str(doc["path"]).lower().endswith(".pdf"):
            tables = extract_tables_from_pdf(doc["path"])
            for t in tables:
                table_text = (
                    f"[TABLE from page {t['page']} — {t['rows']}×{t['cols']}]\n"
                    f"{t['markdown']}"
                )
                chunks.append({
                    "source_name": doc["name"],
                    "chunk_id": chunk_id,
                    "text": table_text,
                    "page": str(t["page"]),
                    "is_table": True,
                })
                chunk_id += 1

    return chunks


# ── Backward-compatible alias ──
chunk_papers = chunk_documents
